from docx import Document
from docx.shared import Inches


def insert_image_into_word(image_path, doc):
    doc.add_picture(image_path, width=Inches(3))  # 调整图片宽度
    doc.add_paragraph()  # 插入一个空段落


def create_word_document(images=None):
    if images is None:
        images = ['plot1.png', 'plot2.png', 'plot3.png', 'plot4.png', 'plot5.png',
                  'plot6.png','plot7.png','plot8.png'
            #, 'plot9.png', 'plot10.png','plot11.png','plot12.png', 'plot13.png'
                  ]
    doc = Document()
    for image in images:
        insert_image_into_word(image, doc)
    doc.save('金融期货图片文档.docx')



