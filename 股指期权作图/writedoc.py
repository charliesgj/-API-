from docx import Document
from docx.shared import Inches
import os

def insert_image_into_word(image_path, doc):
    doc.add_picture(image_path, width=Inches(3))  # 调整图片宽度
    doc.add_paragraph()  # 插入一个空段落


def create_word_document(images=None):
    if images is None:
        images = []
        for j in ["主力合约", "次主力合约"]:
            for k in ["上证50", "中证1000", "沪深300"]:
                images.append(f"{k}{j}希腊字母.png")
        for i in ["上证50", "中证1000", "沪深300"]:
            images.append(f"{i}波动率曲线.png")
        for i in ["上证50", "中证1000", "沪深300"]:
            images.append(f"{i}持仓成交量.png")
            images.append(f"{i}行权价数据统计.png")
        images.append('标的指数历史波动率.png')
    doc = Document()
    for image in images:
        insert_image_into_word(image, doc)
    #for image in images:
    #   os.remove(image)
    doc.save('股指期权图片文档.docx')
